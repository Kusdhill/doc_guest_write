import docx
import zipfile
import sys
import os
import re
import shutil
import subprocess
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xml.dom.minidom


# verifies that a given filename has .docx extension
# if it does not, an error is thrown
def check_extension(filename):
	extension_flag = 0
	extension = ""
	good_extension = "docx"

	for char in filename:
		if extension_flag==0:
			if char==".":
				extension_flag = 1
		else:
			extension+=char

	if extension!=good_extension:
		sys.exit("must pass in .docx files")


# verifies that file exists
# if it does not, an error is thrown
def check_existence(filename):
	if not os.path.isfile(filename):
		sys.exit("file must exist")


# parses file for names of guests
def find_names(doc_object):
	names = []
	
	for i in range(0, len(doc_object.paragraphs)):
		complete_run = ""
		
		for run in doc_object.paragraphs[i].runs:
			if(run.bold):
				complete_run += run.text

		if(verify_name(complete_run)):
			cleaned_name = clean_name(complete_run)
			names.append(cleaned_name)

	mend(names)

	return names


# Verify that identified bold string is actually a name
def verify_name(text):	
	if(not (" " in text)):
		return False
	if(text=="" or len(text)<2):
		return False
	else:
		first_char = text[0]
		last_char = text[-1]
	if(not first_char.isupper()):
		return verify_name(text[1:])
	if(not (last_char.isalpha() or last_char!="," or last_char!=":" or last_char!=" " or last_char!="\"")):
		return False
	else:
		return True


# Checks for fragmented names
def mend(names):
	for i in range(0,len(names)-2):
		if " " not in names[i]:
			join_indeces(names,i,i+1)


# Joins fragmented names
def join_indeces(list_n, left, right):
	first_name = list_n[left]
	last_name  = list_n[right]
	name_string = first_name+" "+last_name

	del list_n[left]
	del list_n[left]
	list_n.insert(left, name_string)


# Cleans name of unnecessary bolded characters
def clean_name(text):
	last_char = text[-1]
	if(not last_char.isalpha()):
		return clean_name(text[0:-1])
	else:
		return text


# put name and text associated with name into dictionary
def copy_text(names, doc):
	name_with_text = {}
	text_list = []

	j = 0
	all_found = False
	for i in range(0,len(doc.paragraphs)):
		line = doc.paragraphs[i]
		text = line.text
		bold = False
		next_name = False
		new_lines = 0
			
		# looking forward
		if(i<len(doc.paragraphs)-1):
			next_line = doc.paragraphs[i+1]
			next_text = next_line.text

		text_list.append(text)

		# if all names have been found and parser is at the end of the doc
		# add text to the dictionary
		if(all_found and i==len(doc.paragraphs)-1):
			name_with_text[names[j]] = text_list

		# if there is a bold run, set bold to true
		if(len(line.runs)>=2):
			for k in range(0,len(line.runs)):
				if(line.runs[k].bold):
					bold = True

		if(j!=0 and re.search('[a-zA-Z]',text)==None):
			if(re.search('[a-zA-Z]',next_text)==None):
				pass
			else:
				name_with_text[names[j-1]] = text_list
				text_list = []

		# if text is bold and it matches a name, increment j (pointer to lines)
		# if end of names list has been reached and name is found then set all_found to true
		if names[j] in text and bold:
			if(j<len(names)-1):
				j+=1
			else:
				all_found = True

	for name in name_with_text:
		print(name)

	return name_with_text


# parses file for images
# changing back to xml implementation
# reading it line by line as string instead of as xml
def parse_images(filename, names):
	print("in parse images\n")

	guest_image_info = {}

	for i in range(0, len(names)):
		guest_image_info[names[i]] = [False,""]


	stripped_filename = filename[0:-5]
	path = "./"+stripped_filename+"_images/word/document.xml"
	output_xml_path = "./"+stripped_filename+"_output.xml"

	raw_xml = xml.dom.minidom.parse(path)
	xml_string = raw_xml.toprettyxml()

	output_xml = open(output_xml_path, "w")
	output_xml.write(xml_string.encode('utf8'))
	output_xml.close()

	text = open(output_xml_path).readlines()

	for j in range(0,len(names)):
		split_name = names[j].split(" ")
		image_found = False
		for i in range(0,len(text)):
			line = text[i]
			if i>0:
				prev_line = text[i-1]
			if "r:embed=" in line:
				image_found = True
				print(line)
			if image_found and "</w:t>" in line and "</w:rPr>" in prev_line:
				print(line)
				image_found = False


# unzip word file to get source xml
def unzip_word(filename):
	stripped_filename = filename[0:-5]
	path = "./"+stripped_filename

	extract_directory = path+"_images"

	zip_ref = zipfile.ZipFile(filename, 'r')
	zip_ref.extractall(extract_directory)
	zip_ref.close()


# get guest images from doc
def get_images(filename):

	image_locations = []
	stripped_filename = filename[0:-5]
	path = "./"+stripped_filename

	extract_directory = path+"_images"

	zip_ref = zipfile.ZipFile(filename, 'r')
	zip_ref.extractall(extract_directory)
	zip_ref.close()
	extract_path = extract_directory+"/word/media"

	for image in os.listdir(extract_path):
		image_locations.append(extract_path+"/"+image)

	image_locations.sort()
	return image_locations


# finds if a name is contained in a string
def contains_name(text, names):
	for name in names:
		if name in text:
			return True

	return False


# if beginning entry doesnt contain name or is empty, entry is cut off list
def clean_entry_list(entry_list, name_list):
	first_entry = entry_list[0]

	if first_entry=="" or not contains_name(first_entry,name_list):
		del entry_list[0]
		return clean_entry_list(entry_list, name_list)
	else:
		return entry_list


# For each name, create a file, dump the text with images, and save the file
def dump_files(filename, names, copied, images):
	print(images)

	path = "./"+filename[0:-5]+"_created_files/"
	all_guest_images = False

	if os.path.exists(path):
		shutil.rmtree(path)

	if(len(names)==len(images)):
		all_guest_images = True

	os.makedirs(path)
	for i in range(0, len(names)):
		entry = copied[names[i]]
		image_added = False
		save_doc = docx.Document()

		entry = clean_entry_list(entry,names)

		for j in range(0,len(entry)):
			# first line
			if(j==0):
				para = save_doc.add_paragraph("")
				run = para.add_run(entry[j])
				run.bold = True

			if(all_guest_images and not image_added):
				save_doc.add_picture(images[i],width=Inches(1.38), height=Inches(1.38))
				image = save_doc.paragraphs[-1]
				image.alignment = WD_ALIGN_PARAGRAPH.CENTER
				image_added = True

			if j!=0 and entry[j]!="":
				save_doc.add_paragraph(entry[j], style = 'List Bullet')

		save_doc.save(path+names[i]+".docx")


# Clean created files
def clean_files(filename):
	stripped_filename = filename[0:-5]
	path = "./"+stripped_filename
	extract_directory = path+"_images"
	output_xml = "./"+stripped_filename+"_output.xml"

	shutil.rmtree(extract_directory)
	os.remove(output_xml)


# Opens results folder
def open_directory(filename):
	directory_path = "./"+filename[0:-5]+"_created_files/"
	subprocess.call(["open", "-R", directory_path])


def main():
	print("checking command line arguments")
	if len(sys.argv)!=2:
		sys.exit("usage: python parse.py filename.docx")
	else:
		print("verifying file extension")
		check_extension(sys.argv[1])
		print("verifying existence")
		check_existence(sys.argv[1])

	filename = sys.argv[1]

	doc = docx.Document(filename)
	print("finding names")
	names = find_names(doc)
	print("copying text")
	names_with_text = copy_text(names, doc)
	print("unzipping word file")
	unzip_word(filename)
	print("getting images")
	parse_images(filename, names)
	guest_images = get_images(filename)
	print("creating files")
	dump_files(filename, names, names_with_text, guest_images)
	print("cleaning created files")
	#clean_files(filename)
	print("opening results")
	#open_directory(filename)


if __name__ == '__main__':
	main()